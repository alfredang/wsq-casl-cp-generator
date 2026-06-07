from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models import ExtractedData

HEADING_COLOR = RGBColor(0x44, 0x72, 0xC4)  # Steel blue matching PDF template
DAY_START_MINUTES = 9 * 60  # 9:00 AM in minutes from midnight
LUNCH_DURATION = 60  # 1 hour lunch
DAY_TOTAL_MINUTES = 480  # 8 hours of usable time (9 AM to 6 PM minus 1h lunch)


def _fmt_time(minutes_from_midnight: int) -> str:
    """Format minutes from midnight as clock time (e.g. 540 -> '9:00', 800 -> '1:20')."""
    h = minutes_from_midnight // 60
    m = minutes_from_midnight % 60
    # Use 12-hour display without AM/PM (matching PDF template style)
    if h > 12:
        h -= 12
    elif h == 0:
        h = 12
    return f"{h}:{m:02d}"


def _build_schedule(data: ExtractedData) -> dict[int, list[dict]]:
    """Build a per-day schedule of time slots from extracted data.

    Returns {day_num: [{"start": str, "end": str, "label": str}, ...]}.
    """
    # Group topics by day
    topics_by_day: dict[int, list] = {}
    for lo in data.learning_outcomes:
        topics_by_day.setdefault(lo.day, []).append(lo)

    # Group assessment by day
    assess_by_day: dict[int, int] = {}
    for am in data.assessment_modes:
        assess_by_day[am.day] = assess_by_day.get(am.day, 0) + am.duration_minutes

    num_days = max(topics_by_day.keys()) if topics_by_day else 1
    schedule: dict[int, list[dict]] = {}

    for day in range(1, num_days + 1):
        slots: list[dict] = []
        topics = topics_by_day.get(day, [])
        assess_min = assess_by_day.get(day, 0)

        # Available instruction time = total day minus assessment
        instruction_min = DAY_TOTAL_MINUTES - assess_min
        num_topics = len(topics)
        per_topic = instruction_min // num_topics if num_topics else 0

        current = DAY_START_MINUTES
        lunch_inserted = False

        for topic in topics:
            # Insert lunch if we've crossed ~1:00 PM and haven't yet
            if not lunch_inserted and current >= 13 * 60 - 10:
                lunch_end = current + LUNCH_DURATION
                slots.append({
                    "start": _fmt_time(current),
                    "end": _fmt_time(lunch_end),
                    "label": "Lunch Break",
                })
                current = lunch_end
                lunch_inserted = True

            end = current + per_topic
            # Strip the topic prefix (e.g. "T1: ...") is already in the topic string
            slots.append({
                "start": _fmt_time(current),
                "end": _fmt_time(end),
                "label": topic.topic,
            })
            current = end

        # Insert lunch after last topic if not yet (unlikely but safe)
        if not lunch_inserted and assess_min > 0:
            lunch_end = current + LUNCH_DURATION
            slots.append({
                "start": _fmt_time(current),
                "end": _fmt_time(lunch_end),
                "label": "Lunch Break",
            })
            current = lunch_end

        # Assessment slot at end of day
        if assess_min > 0:
            end = current + assess_min
            slots.append({
                "start": _fmt_time(current),
                "end": _fmt_time(end),
                "label": "Assessment",
            })

        schedule[day] = slots

    return schedule


def _extract_overview(data: ExtractedData) -> str:
    """Extract a concise course overview paragraph from about_course text."""
    text = data.particulars.about_course
    # Find first substantive paragraph (skip section headers like "a. Benefits...")
    for para in text.split("\n"):
        stripped = para.strip()
        if not stripped:
            continue
        # Skip section header lines (e.g. "a. Benefits of the Course...")
        if len(stripped) < 80:
            continue
        if stripped.startswith("- "):
            continue
        return stripped
    return text[:500]


def _add_colored_heading(doc, text: str, level: int = 2):
    """Add a heading with the steel blue color matching the PDF template."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = HEADING_COLOR


def generate_lesson_plan(data: ExtractedData, output_path: Path) -> Path:
    doc = Document()

    # Configure default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # --- Title ---
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"Lesson Plan: {data.particulars.course_title}")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Calibri"

    # --- Metadata lines ---
    num_days = max(lo.day for lo in data.learning_outcomes) if data.learning_outcomes else 1
    unique_methods = list(dict.fromkeys(im.method for im in data.instruction_methods))

    # Parse hours from summary strings like "14 hour 0 minutes"
    training_hours = data.summary.total_instructional_duration
    assessment_hours = data.summary.total_assessment_duration

    metadata_lines = [
        f"Course Duration: {num_days} Days (9:00 AM \u2013 6:00 PM daily)",
        f"Total Training Hours: {training_hours} (excluding lunch breaks)",
        f"Total Assessment Hours: {assessment_hours}",
        f"Instructional Methods: {', '.join(m.lower() for m in unique_methods)}",
    ]
    for line in metadata_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)

    # --- Course Overview ---
    _add_colored_heading(doc, "Course Overview")
    overview_text = _extract_overview(data)
    doc.add_paragraph(overview_text)

    # --- Day schedules ---
    schedule = _build_schedule(data)

    for day_num in sorted(schedule.keys()):
        _add_colored_heading(doc, f"Day {day_num}")
        for slot in schedule[day_num]:
            slot_text = f"{slot['start']} \u2013 {slot['end']} | {slot['label']}"
            doc.add_paragraph(slot_text)

    doc.save(str(output_path))
    return output_path


def generate_simple_lesson_plan_docx(
    course_title: str,
    days: list[list[dict]],
    output_path: Path,
    topic_minutes: int | None = None,
    course_duration_hrs: float | None = None,
) -> Path:
    """Generate a simple lesson plan .docx with a 4-column table per day:
    Time, Topics, Instructional Methods, Resources.

    ``days`` is the ``days`` list returned by ``build_simple_lesson_plan`` —
    each item is a list of row dicts with keys: time, topic, method, resources.
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"Lesson Plan: {course_title}")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Arial"

    # Metadata
    duration_label = (
        f"{course_duration_hrs:g} hrs / {len(days)} Day(s)"
        if course_duration_hrs
        else f"{len(days)} Day(s)"
    )
    meta = [f"Course Duration: {duration_label} (9:00 AM - 6:00 PM daily)"]
    if topic_minutes:
        meta.append(f"Duration per Topic: {topic_minutes} mins")
    for line in meta:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    headers = ["Time", "Topics", "Instructional Methods", "Resources"]
    col_widths = [Inches(1.2), Inches(2.6), Inches(1.6), Inches(1.1)]
    keys = ["time", "topic", "method", "resources"]

    for day_idx, rows in enumerate(days, start=1):
        _add_colored_heading(doc, f"Day {day_idx}")

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.autofit = False
        for c, w in enumerate(col_widths):
            table.columns[c].width = w

        for c, header in enumerate(headers):
            _set_header_cell(table.rows[0].cells[c], header)

        for row_data in rows:
            row = table.add_row()
            for j, key in enumerate(keys):
                cell = row.cells[j]
                cell.width = col_widths[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(row_data.get(key, "")))
                run.font.name = "Arial"
                run.font.size = Pt(11)

        doc.add_paragraph()

    doc.save(str(output_path))
    return output_path


def _set_header_cell(cell, text: str):
    from docx.oxml.ns import qn
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"), {qn("w:fill"): "4472C4", qn("w:val"): "clear"},
    )
    shading.append(shading_elem)


def generate_lesson_plan_table(
    course_title: str,
    course_duration_hrs: int,
    instructional_hrs: int,
    assessment_hrs: int,
    schedule: dict[int, list[dict]],
    output_path: Path,
    instructional_methods: list[str] | None = None,
) -> Path:
    """Generate a lesson plan .docx with 4-column table (Timing, Duration, Description, Methods)."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"Lesson Plan: {course_title}")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Arial"

    # Metadata
    num_days = len(schedule)
    methods_text = ", ".join(instructional_methods) if instructional_methods else "N/A"
    metadata = [
        f"Course Duration: {course_duration_hrs} hrs / {num_days} Day(s) (9:00 AM - 6:00 PM daily)",
        f"Total Instructional Hours: {instructional_hrs} hrs",
        f"Total Assessment Hours: {assessment_hrs} hrs",
        f"Instructional Methods: {methods_text}",
    ]
    for line in metadata:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    # Day tables
    for day_num in sorted(schedule.keys()):
        _add_colored_heading(doc, f"Day {day_num}")

        rows = schedule[day_num]
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        # Disable autofit so column widths are fixed and text wraps
        table.autofit = False
        # Set column widths to use full page width (6.5" on Letter)
        table.columns[0].width = Inches(1.3)
        table.columns[1].width = Inches(0.9)
        table.columns[2].width = Inches(2.0)
        table.columns[3].width = Inches(2.3)

        _set_header_cell(table.rows[0].cells[0], "Timing")
        _set_header_cell(table.rows[0].cells[1], "Duration")
        _set_header_cell(table.rows[0].cells[2], "Description")
        _set_header_cell(table.rows[0].cells[3], "Instructional Methods")

        col_widths = [Inches(1.3), Inches(0.9), Inches(2.0), Inches(2.3)]
        for row_data in rows:
            row = table.add_row()
            for j, key in enumerate(("timing", "duration", "description", "methods")):
                cell = row.cells[j]
                cell.width = col_widths[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(row_data.get(key, ""))
                run.font.name = "Arial"
                run.font.size = Pt(11)

        doc.add_paragraph()

    doc.save(str(output_path))
    return output_path
