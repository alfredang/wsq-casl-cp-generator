from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models import ExtractedData


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 10):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold


def _add_table_header(table, headers: list[str]):
    """Style header row with bold text and gray background."""
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, text, bold=True, size=10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set header background color
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(
            qn("w:shd"),
            {qn("w:fill"): "4472C4", qn("w:val"): "clear"},
        )
        shading.append(shading_elem)
        # White text for header
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)


def _add_field(doc, label: str, value: str):
    """Add a bold label followed by value text."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_value = p.add_run(str(value))
    run_value.font.size = Pt(11)


def _add_long_text(doc, heading: str, text: str):
    """Add a heading and multi-paragraph text block."""
    doc.add_heading(heading, level=2)
    if not text:
        return
    paragraphs = text.split("\n")
    for para_text in paragraphs:
        stripped = para_text.strip()
        if not stripped:
            continue
        p = doc.add_paragraph(stripped)
        p.style.font.size = Pt(11)


def generate_docx(data: ExtractedData, output_path: Path) -> Path:
    doc = Document()

    # Configure default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading(data.particulars.course_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Training Provider: {data.particulars.training_provider}")
    run.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"Course Type: {data.particulars.course_type}")
    run2.font.size = Pt(12)

    doc.add_page_break()

    # --- Section 1: Course Particulars ---
    doc.add_heading("Section 1: Course Particulars", level=1)
    _add_field(doc, "Name of Registered Training Provider", data.particulars.training_provider)
    _add_field(doc, "Course Title", data.particulars.course_title)
    _add_field(doc, "Course Type", data.particulars.course_type)
    _add_long_text(doc, "About This Course", data.particulars.about_course)
    _add_long_text(doc, "What You Will Learn", data.particulars.what_youll_learn)
    _add_field(doc, "Unique Skill Name", ", ".join(data.particulars.unique_skill_names))

    doc.add_page_break()

    # --- Section 2: Course Background ---
    doc.add_heading("Section 2: Course Background", level=1)
    _add_long_text(doc, "Targeted Sectors and Background", data.background.targeted_sectors)
    if data.background.performance_gaps:
        _add_long_text(doc, "Performance Gaps", data.background.performance_gaps)

    doc.add_page_break()

    # --- Section 3: Instructional Design ---
    doc.add_heading("Section 3: Instructional Design", level=1)

    # Learning Outcomes table
    doc.add_heading("Learning Outcomes", level=2)
    lo_headers = ["Day", "Duration (min)", "LO#", "Learning Outcome", "Topic"]
    table_lo = doc.add_table(rows=1, cols=len(lo_headers))
    table_lo.style = "Table Grid"
    _add_table_header(table_lo, lo_headers)
    for lo in data.learning_outcomes:
        row = table_lo.add_row()
        row.cells[0].text = str(lo.day)
        row.cells[1].text = str(lo.duration_minutes)
        row.cells[2].text = lo.lo_number
        row.cells[3].text = lo.learning_outcome
        row.cells[4].text = lo.topic

    doc.add_paragraph()  # spacing

    # Instruction Methods table
    doc.add_heading("Instruction Methods", level=2)
    meth_headers = ["Day", "Method", "Duration (min)", "Mode of Training"]
    table_meth = doc.add_table(rows=1, cols=len(meth_headers))
    table_meth.style = "Table Grid"
    _add_table_header(table_meth, meth_headers)
    for im in data.instruction_methods:
        row = table_meth.add_row()
        row.cells[0].text = str(im.day)
        row.cells[1].text = im.method
        row.cells[2].text = str(im.duration_minutes)
        row.cells[3].text = im.mode_of_training

    doc.add_page_break()

    # --- Section 4: Assessment ---
    doc.add_heading("Section 4: Assessment", level=1)
    assess_headers = ["Day", "Mode of Assessment", "Duration (min)", "# Assessors", "# Candidates"]
    table_assess = doc.add_table(rows=1, cols=len(assess_headers))
    table_assess.style = "Table Grid"
    _add_table_header(table_assess, assess_headers)
    for am in data.assessment_modes:
        row = table_assess.add_row()
        row.cells[0].text = str(am.day)
        row.cells[1].text = am.mode
        row.cells[2].text = str(am.duration_minutes)
        row.cells[3].text = str(am.num_assessors)
        row.cells[4].text = str(am.num_candidates)

    doc.add_page_break()

    # --- Summary ---
    doc.add_heading("Summary", level=1)

    # (1) Topics
    doc.add_heading("(1) Topics covered in this course", level=2)
    for lo in data.learning_outcomes:
        doc.add_paragraph(lo.topic, style="List Number")

    # (2) Instructional methods
    doc.add_heading("(2) Instructional methods", level=2)
    unique_methods = list(dict.fromkeys(im.method for im in data.instruction_methods))
    doc.add_paragraph(", ".join(unique_methods))

    # (3) Duration for each topic
    doc.add_heading("(3) Duration for each topic", level=2)
    dur_headers = ["Topic", "Duration (min)"]
    table_dur = doc.add_table(rows=1, cols=len(dur_headers))
    table_dur.style = "Table Grid"
    _add_table_header(table_dur, dur_headers)
    for lo in data.learning_outcomes:
        row = table_dur.add_row()
        row.cells[0].text = lo.topic
        row.cells[1].text = str(lo.duration_minutes)

    doc.add_paragraph()  # spacing

    # Course totals
    doc.add_heading("Course Totals", level=2)
    _add_field(doc, "Total Course Duration", data.summary.total_course_duration)
    _add_field(doc, "Total Instructional Duration", data.summary.total_instructional_duration)
    _add_field(doc, "Total Assessment Duration", data.summary.total_assessment_duration)
    _add_field(doc, "Mode of Training", data.summary.mode_of_training)

    doc.save(str(output_path))
    return output_path
